import logging
from telegram import Bot as TelegramBot

from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup, ForceReply
from telegram.constants import ParseMode
from telegram.ext import (
    ContextTypes,
    ConversationHandler,
    CallbackQueryHandler,
    CommandHandler,
    MessageHandler,
    filters,
)

from database.queries import DatabaseQueries
import config
from utils.helpers import is_admin
from utils.helpers import staff_only_decorator as staff_only

import json
from ai.model import responder
import html  # For escaping HTML entities
from io import BytesIO
from docx import Document
from docx.shared import RGBColor
from telegram.constants import ParseMode
from telegram.helpers import escape_markdown

logger = logging.getLogger(__name__)

class AdminTicketHandler:
    """Handle ticket management for admins"""

    # ---------------------------------------------------------------------
    # Helper methods
    # ---------------------------------------------------------------------
    def _is_admin(self, telegram_id: int) -> bool:
        """Return True if the given Telegram user id is an *admin* OR *support* user.

        This logic mirrors the `admin_only_decorator` so that callback / internal
        code paths can quickly verify permissions without applying the decorator
        again. We purposefully allow *support* users (stored in the `support_users`
        table) to access the handler in addition to configured admins.
        """
        from utils.helpers import is_user_in_admin_list
        is_admin_flag = is_user_in_admin_list(telegram_id, self.admin_config)
        is_support_flag = DatabaseQueries.is_support_user(telegram_id)
        return bool(is_admin_flag or is_support_flag)

    # --------------------------- Ticket look-ups --------------------------
    def _get_all_tickets(self):
        """Return *all* tickets as a list of dicts (newest first)."""
        return DatabaseQueries.get_all_tickets()

    def _get_pending_tickets(self):
        """Return open / pending tickets that need admin attention."""
        # Prefer dedicated helper if available; fall back to generic open tickets.
        if hasattr(DatabaseQueries, "get_open_tickets"):
            return DatabaseQueries.get_open_tickets()
        # Fallback – select tickets with status not closed.
        tickets = DatabaseQueries.get_all_tickets()
        return [t for t in tickets if (t.get("status") or "").lower() != "closed"]

    def _get_ticket_by_id(self, ticket_id: int):
        """Fetch a single ticket row (dict) by its id, or None."""
        row = DatabaseQueries.get_ticket(ticket_id)
        if row and not isinstance(row, dict):
            # Convert sqlite3.Row or tuple using column description if needed
            try:
                # Attempt to construct from row keys (sqlite3.Row is mapping-like)
                row = dict(row)
            except Exception:
                pass
        return row

    def _get_original_question(self, ticket_id: int) -> str:
        """Return the first (original) message text of a ticket, if exists."""
        try:
            messages = DatabaseQueries.get_ticket_messages(ticket_id)
            if not messages:
                return ""
            first = messages[0]
            if not isinstance(first, dict):
                try:
                    first = dict(first)
                except Exception:
                    pass
            return first.get("message", "") or ""
        except Exception as e:
            logger.error(f"Error fetching original ticket question for {ticket_id}: {e}")
            return ""

    # --------------------------- User helpers -----------------------------
    def _get_user_info(self, user_id: int):
        """Return minimal user info dict for display (full_name / username)."""
        if user_id is None:
            return {}
        info = DatabaseQueries.get_user_by_telegram_id(user_id)
        if info and not isinstance(info, dict):
            try:
                info = dict(info)
            except Exception:
                pass
        return info or {}

    def _format_user_info(self, user_info: dict) -> str:
        """Return a human-readable representation of a user for ticket lists."""
        if not user_info:
            return "نامشخص"
        name = user_info.get("full_name") or ""
        username = user_info.get("username") or ""
        if username:
            username = f"@{username}"
        if name and username:
            return f"{name} ({username})"
        return name or username or str(user_info.get("user_id", "نامشخص"))

    # -----------------------------------------------------------------
    # Contact info helper
    # -----------------------------------------------------------------

    def _get_contact_info(self, user_info: dict) -> str:
        """Return contact information string for a user (username / phone)."""
        if not user_info:
            return "نامشخص"
        parts = []
        username = user_info.get("username")
        if username:
            parts.append(f"@{username}")
        phone = user_info.get("phone") or user_info.get("phone_number")
        if phone:
            parts.append(phone)
        email = user_info.get("email")
        if email:
            parts.append(email)
        return " | ".join(parts) if parts else "نامشخص"

    # -----------------------------------------------------------------
    # Ticket messages helpers
    # -----------------------------------------------------------------

    def _get_original_question(self, ticket_id: int) -> str:
        """Return the first *user* message for given ticket (original question)."""
        try:
            msgs = DatabaseQueries.get_ticket_messages(ticket_id)
            if not msgs:
                return "-"
            # Find first non-admin message
            for m in msgs:
                msg_dict = dict(m)
                if not msg_dict.get("is_admin"):
                    return msg_dict.get("message", "-") or "-"
            # Fallback to first message
            return dict(msgs[0]).get("message", "-")
        except Exception:
            return "-"

    def _get_admin_reply(self, ticket_id: int) -> str | None:
        """Return the latest admin/support reply for the given ticket, or None if not found."""
        try:
            msgs = DatabaseQueries.get_ticket_messages(ticket_id)
            if not msgs:
                return None
            # Iterate from newest to oldest to get last admin message
            for m in reversed(msgs):
                msg_dict = dict(m)
                if msg_dict.get("is_admin"):
                    raw = msg_dict.get("message")
                    if isinstance(raw, dict):
                        return raw.get("text", str(raw))
                    else:
                        return str(raw) if raw is not None else None
            return None
        except Exception:
            return None

        # -----------------------------------------------------------------
    # UI origin helper
    # -----------------------------------------------------------------
    def _determine_origin_list(self, query) -> str:
        """Return 'all' if the current inline message originates from *all tickets* list, else 'pending'."""
        try:
            kbd = query.message.reply_markup
            if not kbd:
                return 'pending'
            for row in kbd.inline_keyboard:
                for btn in row:
                    cb = getattr(btn, 'callback_data', '') or ''
                    if cb.startswith('all_tickets_page_') or cb == 'refresh_all_tickets':
                        return 'all'
        except Exception:
            pass
        return 'pending'

    # --------------------------- Misc helpers -----------------------------
    _STATUS_EMOJI_MAP = {
        "open": "🟢",
        "pending_admin_reply": "🟡",
        "pending_user_reply": "🟠",
        "closed": "🔴",
    }

    def _get_status_emoji(self, status: str) -> str:
        """Map internal ticket status to an emoji used in UIs."""
        return self._STATUS_EMOJI_MAP.get(str(status).lower(), "❔")
    
    def __init__(self):
        """Initialize the handler and set `admin_config` required by permission decorators."""
        self.admin_config = getattr(config, "MANAGER_BOT_ADMINS_DICT", {}) or getattr(config, "ADMIN_USER_IDS", [])
        # Use the *main* bot token to reach end-users; manager bot cannot initiate chats.
        try:
            self.main_bot = TelegramBot(token=config.MAIN_BOT_TOKEN)
        except Exception as e:  # Fallback if token missing
            logger.error("Failed to create main bot instance: %s", e)
            self.main_bot = None
    
    @staff_only
    async def show_ticket_history_for_user(self, update: Update, context: ContextTypes.DEFAULT_TYPE, target_user_id: int):
        """Show recent tickets of specified user to admin/support."""
        tickets = DatabaseQueries.get_tickets_by_user(target_user_id, limit=20)
        if not tickets:
            await update.effective_message.reply_text("📭 برای این کاربر تیکتی یافت نشد.")
            return
        text_lines = [f"📄 *آخرین تیکت‌های کاربر {target_user_id}:*\n"]
        keyboard = []
        for t in tickets:
            t = dict(t)
            ticket_id = t.get('id') or t.get('ticket_id')
            subject = t.get('subject', 'بدون موضوع')
            status = (t.get('status') or '').replace('_', ' ')
            created_at = t.get('created_at', '')
            text_lines.append(f"• #{ticket_id} | {subject} | {status} | {created_at}")
            keyboard.append([InlineKeyboardButton(f"تیکت#{ticket_id}", callback_data=f"view_ticket_{ticket_id}")])
        await update.effective_message.reply_text("\n".join(text_lines), parse_mode=ParseMode.MARKDOWN, reply_markup=InlineKeyboardMarkup(keyboard))

    def get_ticket_conversation_handler(self):
        """Get the conversation handler for ticket management"""
        return ConversationHandler(
            entry_points=[
                CallbackQueryHandler(self.show_tickets_command, pattern="^show_tickets$"),
                CallbackQueryHandler(self.view_ticket_callback, pattern="^view_ticket_\\d+$"),
                CallbackQueryHandler(self.generate_answer_callback, pattern="^gen_answer_\\d+$"),
                CallbackQueryHandler(self.manual_answer_callback, pattern="^manual_answer_\\d+$"),
                CallbackQueryHandler(self.edit_answer_callback, pattern="^edit_answer_\\d+$"),
                CallbackQueryHandler(self.send_answer_callback, pattern="^send_answer_\\d+$"),
                CallbackQueryHandler(self.close_ticket_callback, pattern="^close_ticket_\\d+$"),
                CallbackQueryHandler(self.paginate_tickets_callback, pattern=r'^tickets_page_\d+$'),
                CallbackQueryHandler(lambda u, c: u.callback_query.answer(), pattern=r'^ignore$'),
                CallbackQueryHandler(self.refresh_tickets_callback, pattern="^refresh_tickets$"),
                CallbackQueryHandler(self.refresh_all_tickets_callback, pattern="^refresh_all_tickets$"),
                CallbackQueryHandler(self.export_all_tickets_callback, pattern="^export_all_tickets$"),
                CallbackQueryHandler(self.paginate_all_tickets_callback, pattern=r'^all_tickets_page_\d+$')
            ],
            states={
                # Add states as needed
            },
            fallbacks=[
                CallbackQueryHandler(self.show_tickets_command, pattern="^refresh_tickets$")
            ],
            per_message=False
        )
    
    @staff_only
    async def show_tickets_command(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Show all pending tickets to admin"""
        user_id = update.effective_user.id
        
        try:
            # Get all pending tickets
            tickets = self._get_pending_tickets()
            
            if not tickets:
                await update.message.reply_text("📋 هیچ تیکت باز یافت نشد.")
                return
            
            # Create ticket list with inline keyboard
            keyboard = []
            row = []
            message_text = "📋 لیست تیکت‌های باز:\n\n"
            
            per_page = 10
            page = 0
            start = page * per_page
            end = start + per_page
            tickets_page = tickets[start:end]
            
            for ticket in tickets_page:  # Show max 10 tickets at once
                ticket = dict(ticket)  # تبدیل Row به دیکشنری
                ticket_id = ticket.get('ticket_id') or ticket.get('id')
                user_id_ticket = ticket.get('user_id')
                subject = ticket.get('subject', 'بدون موضوع')
                created_at = ticket.get('created_at', 'نامشخص')
                
                # Get user info
                user_info = self._get_user_info(user_id_ticket)
                user_display = self._format_user_info(user_info)
                
                message_text += f"تیکت #{ticket_id}\n"
                message_text += f"کاربر: {user_display}\n"
                message_text += f"موضوع: {subject}\n"
                message_text += f"تاریخ: {created_at}\n"
                message_text += "───────────────────\n\n"
                
                # Add button for this ticket to the current row
                row.append(
                    InlineKeyboardButton(
                        f"تیکت#{ticket_id}", 
                        callback_data=f"view_ticket_{ticket_id}"
                    )
                )
                
                # If row is full (3 items), add it to the keyboard and start a new row
                if len(row) == 3:
                    keyboard.append(row)
                    row = []
            
            # Add any remaining buttons in the last row (if it's not empty)
            if row:
                keyboard.append(row)

            # Navigation row
            nav_row = []
            if end < len(tickets):
                nav_row.append(InlineKeyboardButton("بعدی »", callback_data="tickets_page_1"))
            if nav_row:
                keyboard.append(nav_row)

            # Add refresh button on its own row
            keyboard.append([
                InlineKeyboardButton("به‌روزرسانی لیست تیکت‌ها", callback_data="refresh_tickets")
            ])

            # Add export tickets button on its own row
            keyboard.append([
                InlineKeyboardButton("📄 خروجی تیکت‌ها", callback_data="export_all_tickets")
            ])
            
            reply_markup = InlineKeyboardMarkup(keyboard)
            await update.message.reply_text(
                message_text,
                reply_markup=reply_markup
            )
            
        except Exception as e:
            logger.error(f"Error showing tickets: {e}")
            await update.message.reply_text("خطا در نمایش تیکت‌ها.")

    async def view_ticket_callback(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Show detailed view of a specific ticket"""
        query = update.callback_query
        await query.answer()

        user_id = query.from_user.id

        # Check if user is admin
        if not self._is_admin(user_id):
            await query.edit_message_text("شما دسترسی لازم برای این عملیات را ندارید.")
            return

        try:
            # Extract ticket ID from callback data
            ticket_id = int(query.data.split('_')[-1])
            logger.info(f"DEBUG: Admin {user_id} viewing ticket with ID: {ticket_id}")

            # Get ticket details
            ticket = self._get_ticket_by_id(ticket_id)

            if not ticket:
                logger.warning(f"Ticket with ID {ticket_id} not found by _get_ticket_by_id.")
                await query.edit_message_text("تیکت یافت نشد.")
                return

            logger.info(f"DEBUG: Ticket data for display: {ticket}")

            user_id_ticket = ticket.get('user_id')
            subject = ticket.get('subject', 'بدون موضوع')
            message_raw = ticket.get('message') or self._get_original_question(ticket_id)
            # Ensure message is a plain string for html.escape
            if isinstance(message_raw, dict):
                # Try common textual keys first
                for key in ("text", "body", "message", "content"):
                    if key in message_raw and isinstance(message_raw[key], str):
                        message = message_raw[key]
                        break
                else:
                    # Fallback: dump JSON for readability
                    import json as _json
                    message = _json.dumps(message_raw, ensure_ascii=False)
            else:
                message = str(message_raw)
            # Final guard: ensure message is str
            if not isinstance(message, str):
                message = str(message)
            created_at = ticket.get('created_at', 'نامشخص')
            status = ticket.get('status', 'نامشخص')



            # Get user info
            user_info = self._get_user_info(user_id_ticket)
            if not user_info:
                await query.edit_message_text("اطلاعات کاربر یافت نشد.")
                return


            user_display = html.escape(self._format_user_info(user_info))
            contact_info = html.escape(self._get_contact_info(user_info))
            subject_html = html.escape(subject)
            message_html = html.escape(message)

            # Retrieve last admin reply if exists
            admin_reply = self._get_admin_reply(ticket_id)

            admin_reply_html = html.escape(admin_reply) if admin_reply else None

            # Format ticket details
            message_text = (
                f"جزئیات تیکت #{ticket_id}\n\n"
                f"کاربر: {user_display}\n"
                f"موضوع: {subject_html}\n"
                f"تاریخ ایجاد: {created_at}\n"
                f"وضعیت: {self._get_status_emoji(status)} {status}\n\n"
                f"پیام:\n{message_html}\n"
            )
            if admin_reply_html:
                message_text += f"\nپاسخ پشتیبانی:\n{admin_reply_html}\n"
            message_text += f"\nاطلاعات تماس: {contact_info}"



                        # Determine which list we should return to (all vs pending)
            back_cb = "refresh_all_tickets" if self._determine_origin_list(query) == "all" else "refresh_tickets"
            # Create action buttons
            keyboard = [
                [  # First row: generate AI answer
                    InlineKeyboardButton("🔄 تولید پاسخ پیشنهادی", callback_data=f"gen_answer_{ticket_id}"),
                    InlineKeyboardButton("✍️ پاسخ دستی", callback_data=f"manual_answer_{ticket_id}")
                ],
                [  # Second row: Close ticket
                    InlineKeyboardButton("بستن تیکت", callback_data=f"close_ticket_{ticket_id}")
                ],
                [  # Third row: Back to list
                    InlineKeyboardButton("بازگشت به لیست تیکت‌ها", callback_data=back_cb)
                ]
            ]

            reply_markup = InlineKeyboardMarkup(keyboard)
            await query.edit_message_text(
                message_text,
                reply_markup=reply_markup
            )

        except Exception as e:
            logger.error(f"Error viewing ticket: {e}")
            await query.edit_message_text("خطا در نمایش تیکت.")

    def _determine_origin_list(self, query) -> str:
        """Return 'all' if ticket view originated from the all-tickets list, else 'pending'.
        We inspect the text of the message containing the list.
        """
        try:
            msg_text = query.message.text or ""
            if "تمام تیکت‌ها" in msg_text:
                return "all"
        except Exception:
            pass
        return "pending"

    async def generate_answer_callback(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Generate AI suggested answer on demand and update the ticket view"""
        query = update.callback_query
        await query.answer("در حال تولید پاسخ...")

        user_id = query.from_user.id
        if not self._is_admin(user_id):
            await query.edit_message_text("شما دسترسی لازم برای این عملیات را ندارید.")
            return

        ticket_id = int(query.data.split('_')[-1])
        ticket = self._get_ticket_by_id(ticket_id)
        if not ticket:
            await query.edit_message_text("تیکت یافت نشد.")
            return

        # Generate AI answer
        subject = ticket.get('subject', 'بدون موضوع')
        first_message = ticket.get('message', '')
        ai_answer = responder.answer_ticket(subject, first_message, ticket.get('user_id'))
        if not ai_answer:
            await query.edit_message_text("خطا در تولید پاسخ توسط هوش‌مصنوعی.")
            return
        context.user_data[f'ai_answer_{ticket_id}'] = ai_answer

        # Prepare updated message text (similar to view_ticket details)
        user_info = self._get_user_info(ticket.get('user_id'))
        user_display = html.escape(self._format_user_info(user_info)) if user_info else "نامشخص"
        contact_info = html.escape(self._get_contact_info(user_info)) if user_info else "-"
        subject_html = html.escape(subject)
        message_html = html.escape(first_message)
        created_at = ticket.get('created_at', 'نامشخص')
        status = ticket.get('status', 'نامشخص')

        message_text = (
            f"جزئیات تیکت #{ticket_id}\n\n"
            f"کاربر: {user_display}\n"
            f"موضوع: {subject_html}\n"
            f"تاریخ ایجاد: {created_at}\n"
            f"وضعیت: {self._get_status_emoji(status)} {status}\n\n"
            f"پیام:\n{message_html}\n\n"
            f"اطلاعات تماس: {contact_info}\n\n"
            f"پاسخ پیشنهادی هوش‌مصنوعی:\n{ai_answer}\n"
        )

        # Updated keyboard with send/edit options
        keyboard = [
            [
                InlineKeyboardButton("ارسال پاسخ پیشنهادی", callback_data=f"send_answer_{ticket_id}"),
                InlineKeyboardButton("ویرایش و ارسال پاسخ", callback_data=f"edit_answer_{ticket_id}")
            ],
            [
                InlineKeyboardButton("🔄 تولید مجدد پاسخ", callback_data=f"gen_answer_{ticket_id}")
            ],
            [
                InlineKeyboardButton("بستن تیکت", callback_data=f"close_ticket_{ticket_id}")
            ],
            [
                InlineKeyboardButton("بازگشت به لیست تیکت‌ها", callback_data="refresh_tickets")
            ]
        ]
        await query.edit_message_text(message_text, reply_markup=InlineKeyboardMarkup(keyboard))

    async def send_answer_callback(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Send the stored AI answer to the ticket owner and refresh admin menu so the keyboard remains available."""
        query = update.callback_query
        await query.answer()

        admin_id = query.from_user.id
        if not self._is_admin(admin_id):
            await query.edit_message_text("شما دسترسی لازم برای این عملیات را ندارید.")
            return

        # Extract ticket_id from callback_data e.g. 'send_answer_123'
        try:
            ticket_id = int(query.data.split('_')[-1])
        except (IndexError, ValueError):
            await query.edit_message_text("شناسهٔ تیکت نامعتبر است.")
            return

        ticket = self._get_ticket_by_id(ticket_id)
        if not ticket:
            await query.edit_message_text("تیکت یافت نشد.")
            return

        # Retrieve previously generated AI answer from user_data
        ai_answer = context.user_data.get(f'ai_answer_{ticket_id}')
        if not ai_answer:
            await query.edit_message_text("پاسخ هوش‌مصنوعی یافت نشد.")
            return

        try:
            # Compose text for the ticket owner
            original_question = ticket.get('message') or self._get_original_question(ticket_id)
            user_reply_text = (
                f"❔ سوال شما:\n{original_question}\n\n"
                f"✅ پاسخ:\n{ai_answer}"
            )

            # Send the answer using main bot if available (for multi-bot setups)
            bot_to_use = self.main_bot or context.bot
            await bot_to_use.send_message(chat_id=ticket.get('user_id'), text=user_reply_text)

            # Update DB: store answer as admin message + close ticket
            DatabaseQueries.add_ticket_message(ticket_id, admin_id, ai_answer, is_admin_message=True, update_status=False)
            DatabaseQueries.update_ticket_status(ticket_id, 'closed')

            # Refresh the tickets list so inline keyboard stays accessible
            await self._show_tickets_inline(query, page=0)
        except Exception as e:
            logger.error(f"Error sending AI answer: {e}")
            await query.edit_message_text("خطا در ارسال پاسخ به کاربر.")



    async def edit_answer_callback(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Let admin edit the AI-generated answer before sending."""
        query = update.callback_query
        await query.answer()
        admin_id = query.from_user.id
        if not self._is_admin(admin_id):
            await query.edit_message_text("شما دسترسی لازم برای این عملیات را ندارید.")
            return
        try:
            ticket_id = int(query.data.split('_')[-1])
        except (IndexError, ValueError):
            await query.edit_message_text("شناسهٔ تیکت نامعتبر است.")
            return
        ai_answer = context.user_data.get(f'ai_answer_{ticket_id}')
        if not ai_answer:
            await query.edit_message_text("پاسخ هوش‌مصنوعی یافت نشد.")
            return
        # Send ForceReply with pre-filled answer
        await query.message.reply_text(
            f"لطفاً پاسخ را ویرایش کرده و ارسال کنید. پس از ارسال، ربات آن را به کاربر فوروارد خواهد کرد.\n\nمتن پیشنهادی:\n{ai_answer}",
            reply_markup=ForceReply(selective=True)
        )
        context.user_data['editing_ticket_id'] = ticket_id

    async def manual_answer_callback(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Prompt admin to type a manual answer without AI suggestion"""
        query = update.callback_query
        await query.answer()

        user_id = query.from_user.id
        if not self._is_admin(user_id):
            await query.edit_message_text("شما دسترسی لازم برای این عملیات را ندارید.")
            return

        ticket_id = int(query.data.split('_')[-1])
        ticket = self._get_ticket_by_id(ticket_id)
        if not ticket:
            await query.edit_message_text("تیکت یافت نشد.")
            return

        # Ask admin to reply manually
        await query.message.reply_text(
            "لطفاً پاسخ خود را بنویسید و ارسال کنید.",
            reply_markup=ForceReply(selective=True)
        )
        context.user_data['editing_ticket_id'] = ticket_id



    async def receive_edited_answer(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Handle admin's edited answer after ForceReply"""
        user_id = update.effective_user.id
        if not self._is_admin(user_id):
            return  # ignore non-admins
        ticket_id = context.user_data.pop('editing_ticket_id', None)
        if not ticket_id:
            return  # nothing to process
        ticket = self._get_ticket_by_id(ticket_id)
        if not ticket:
            await update.message.reply_text("تیکت یافت نشد.")
            return
        target_user_id = ticket.get('user_id')
        text = update.message.text
        try:
            original_question = ticket.get('message') or self._get_original_question(ticket_id)
            combined_text = (
                f"❔ سوال شما:\n{original_question}\n\n"
                f"✅ پاسخ:\n{text}"
            )
            bot_to_use = self.main_bot or context.bot
            await bot_to_use.send_message(chat_id=target_user_id, text=combined_text)
            DatabaseQueries.add_ticket_message(ticket_id, user_id, text, is_admin_message=True, update_status=False)
            DatabaseQueries.update_ticket_status(ticket_id, 'closed')
            await update.message.reply_text(
                "💠 پاسخ برای کاربر ارسال شد و تیکت بسته شد.\n\n"
                f"❔ سوال:\n{original_question}\n\n✅ پاسخ ارسال شده:\n{text}"
            )
        except Exception as e:
            logger.error(f"Error forwarding edited answer: {e}")
            await update.message.reply_text("خطا در ارسال پاسخ.")

    async def close_ticket_callback(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Close a ticket"""
        query = update.callback_query
        await query.answer()
        
        user_id = query.from_user.id
        
        # Check if user is admin
        if not self._is_admin(user_id):
            await query.edit_message_text("شما دسترسی لازم برای این عملیات را ندارید.")
            return
        
        try:
            # Extract ticket ID from callback data
            ticket_id = int(query.data.split('_')[-1])
            
            # Close the ticket
            success = self._close_ticket(ticket_id, query.from_user.id) # Pass admin_id for logging or future use
            
            if success:
                # Get ticket info to notify user
                ticket = self._get_ticket_by_id(ticket_id)
                if ticket:
                    ticket_user_id = ticket.get('user_id')
                    
                    # Try to notify the user
                    try:
                        await context.bot.send_message(
                            chat_id=ticket_user_id,
                            text=f"⭕ تیکت شما با شماره #{ticket_id} بسته شد.\n"
                                 f"پاسخ مربوطه از طریق پشتیبانی ارسال خواهد شد."
                        )
                    except Exception:
                        logger.warning(f"Could not notify user {ticket_user_id} about closed ticket")
                
                await query.edit_message_text(
                    f"تیکت #{ticket_id} با موفقیت بسته شد.\n"
                    f"کاربر در صورت امکان از بسته شدن تیکت مطلع شد."
                )
            else:
                await query.edit_message_text("خطا در بستن تیکت.")
                
        except Exception as e:
            logger.error(f"Error closing ticket: {e}")
            await query.edit_message_text("خطا در بستن تیکت.")

    async def refresh_tickets_callback(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Refresh the list of pending tickets (inline)."""
        query = update.callback_query
        # Show a short answering toast but don't block on it
        await query.answer("🔄 در حال به‌روزرسانی…", show_alert=False)

        # Only admins/support are allowed – reuse the same check
        if not self._is_admin(query.from_user.id):
            await query.edit_message_text("❌ شما دسترسی لازم برای این عملیات را ندارید.")
            return

        # Delegate to the inline implementation that already handles
        # the 'Message is not modified' BadRequest gracefully.
        await self._show_tickets_inline(query, page=0)

    async def paginate_tickets_callback(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Paginate tickets"""
        query = update.callback_query
        await query.answer()

        page = int(query.data.split('_')[-1])
        await self._show_tickets_inline(query, page)

    def get_handlers(self):
        """Get all telegram.ext handlers required for ticket management."""
        return [
            # Commands
            CommandHandler('tickets', self.show_tickets_command),
            CommandHandler('all_tickets', self.show_all_tickets_command),

            # Callback queries – common actions
            CallbackQueryHandler(self.view_ticket_callback, pattern=r'^view_ticket_\d+$'),
            CallbackQueryHandler(self.generate_answer_callback, pattern=r'^gen_answer_\d+$'),
            CallbackQueryHandler(self.manual_answer_callback, pattern=r'^manual_answer_\d+$'),
            CallbackQueryHandler(self.send_answer_callback, pattern=r'^send_answer_\d+$'),
            CallbackQueryHandler(self.edit_answer_callback, pattern=r'^edit_answer_\d+$'),
            CallbackQueryHandler(self.close_ticket_callback, pattern=r'^close_ticket_\d+$'),

            # Pending-tickets list callbacks
            CallbackQueryHandler(self.refresh_tickets_callback, pattern=r'^refresh_tickets$'),
            CallbackQueryHandler(self.paginate_tickets_callback, pattern=r'^tickets_page_\d+$'),

            # All-tickets list callbacks
            CallbackQueryHandler(self.refresh_all_tickets_callback, pattern=r'^refresh_all_tickets$'),
        CallbackQueryHandler(self.export_all_tickets_callback, pattern=r'^export_all_tickets$'),
            CallbackQueryHandler(self.paginate_all_tickets_callback, pattern=r'^all_tickets_page_\d+$'),

            # Edited answers (plain text reply)
            MessageHandler(filters.REPLY & filters.TEXT & (~filters.COMMAND), self.receive_edited_answer),
        ]

    async def _show_tickets_inline(self, query, page: int = 0):
        """Show tickets in inline message"""
        try:
            # Get all pending tickets and convert rows to dicts so we can access with .get
            tickets_raw = self._get_pending_tickets()
            tickets = [dict(t) for t in tickets_raw]
            # Sort descending by ticket id (newest first) to keep ordering stable.
            tickets.sort(key=lambda t: (t.get("ticket_id") or t.get("id") or 0), reverse=True)

            if not tickets:
                await query.edit_message_text("هیچ تیکت باز یافت نشد.")
                return

            # Pagination bookkeeping
            logger.debug(f"_show_tickets_inline: total_tickets={len(tickets)}, requested_page={page}")
            per_page = 10
            total_pages = (len(tickets) - 1) // per_page + 1
            # Clamp page index to valid range to avoid empty pages if ticket count changed
            page = max(0, min(page, total_pages - 1))

            # Create ticket list with inline keyboard
            keyboard = []
            message_text = "📋 لیست تیکت‌های باز:\n\n"
            
            start = page * per_page
            end = start + per_page
            tickets_page = tickets[start:end]
            
            row = []
            for ticket in tickets_page:  # Show max 10 tickets at once
                ticket = dict(ticket)  # تبدیل Row به دیکشنری
                ticket_id = ticket.get('ticket_id') or ticket.get('id')
                user_id_ticket = ticket.get('user_id')
                subject = ticket.get('subject', 'بدون موضوع')
                created_at = ticket.get('created_at', 'نامشخص')
                
                # Get user info
                user_info = self._get_user_info(user_id_ticket)
                user_display = self._format_user_info(user_info)
                
                message_text += f"تیکت #{ticket_id}\n"
                message_text += f"کاربر: {user_display}\n"
                message_text += f"موضوع: {subject}\n"
                message_text += f"تاریخ: {created_at}\n"
                message_text += "───────────────────\n\n"
                
                # Add button to current row
                row.append(
                    InlineKeyboardButton(
                        f"تیکت#{ticket_id}",
                        callback_data=f"view_ticket_{ticket_id}"
                    )
                )
                # If row full (3 buttons), push to keyboard
                if len(row) == 3:
                    keyboard.append(row)
                    row = []
            # leftover row
            if row:
                keyboard.append(row)
            
            # Navigation row with current page indicator
            nav_row = []
            if page > 0:
                nav_row.append(InlineKeyboardButton("« قبلی", callback_data=f"tickets_page_{page-1}"))
            nav_row.append(InlineKeyboardButton(f"صفحه {page+1}/{total_pages}", callback_data="ignore"))
            if page < total_pages - 1:
                nav_row.append(InlineKeyboardButton("بعدی »", callback_data=f"tickets_page_{page+1}"))
            if nav_row:
                keyboard.append(nav_row)

            # Add refresh button
            keyboard.append([
                InlineKeyboardButton("به‌روزرسانی", callback_data="refresh_tickets")
            ])

            # Add export tickets button
            keyboard.append([
                InlineKeyboardButton("📄 خروجی تیکت‌ها", callback_data="export_all_tickets")
            ])
            
            reply_markup = InlineKeyboardMarkup(keyboard)
            await query.edit_message_text(
                message_text,
                reply_markup=reply_markup
            )
            
        except Exception as e:
            import telegram
            if isinstance(e, telegram.error.BadRequest) and "Message is not modified" in str(e):
                logger.warning(f"Suppressed Telegram BadRequest: {e}")
                return
            logger.error(f"Error showing tickets inline: {e}")
            await query.edit_message_text("❌ خطا در نمایش تیکت‌ها.")

    async def _show_all_tickets_inline(self, query, page: int = 0):
        """Show all tickets in inline message"""
        try:
            # Get all tickets and convert rows to dicts so we can access with .get
            tickets_raw = self._get_all_tickets()
            tickets = [dict(t) for t in tickets_raw]
            # Sort descending by ticket id (newest first) to keep ordering stable.
            tickets.sort(key=lambda t: (t.get("ticket_id") or t.get("id") or 0), reverse=True)

            if not tickets:
                await query.edit_message_text("هیچ تیکتی یافت نشد.")
                return

            # Pagination bookkeeping
            logger.debug(f"_show_all_tickets_inline: total_tickets={len(tickets)}, requested_page={page}")
            per_page = 10
            total_pages = (len(tickets) - 1) // per_page + 1
            # Clamp page index to valid range to avoid empty pages if ticket count changed
            page = max(0, min(page, total_pages - 1))

            # Create ticket list with inline keyboard
            keyboard = []
            message_text = f"📋 لیست تمام تیکت‌ها (صفحه {page+1}/{total_pages}):\n\n"
            
            start = page * per_page
            end = start + per_page
            tickets_page = tickets[start:end]
            
            row = []
            for ticket in tickets_page:  # Show max 10 tickets at once
                ticket = dict(ticket)  # تبدیل Row به دیکشنری
                ticket_id = ticket.get('ticket_id') or ticket.get('id')
                user_id_ticket = ticket.get('user_id')
                subject = ticket.get('subject', 'بدون موضوع')
                created_at = ticket.get('created_at', '')
                status = ticket.get('status', '')
                readable_status = escape_markdown(str(status).replace('_', ' '), version=1)
                emoji = self._get_status_emoji(status)
                
                message_text += f"{emoji} *تیکت #{ticket_id}* ({readable_status})\n"
                user_display = escape_markdown(self._format_user_info(self._get_user_info(user_id_ticket)), version=1)
                message_text += f"👤 کاربر: {user_display}\n"
                message_text += f"📝 موضوع: {escape_markdown(subject, version=1)}\n"
                message_text += f"📅 تاریخ: {escape_markdown(created_at, version=1)}\n"
                message_text += "───────────────────\n\n"
                
                # Add button to current row
                row.append(
                    InlineKeyboardButton(
                        f"{emoji} #{ticket_id}", callback_data=f"view_ticket_{ticket_id}"
                    )
                )
                # If row full (3 buttons), push to keyboard
                if len(row) == 3:
                    keyboard.append(row)
                    row = []
            # leftover row
            if row:
                keyboard.append(row)
            
            # Navigation row with current page indicator
            nav_row = []
            if page > 0:
                nav_row.append(InlineKeyboardButton("◀️ صفحه قبل", callback_data=f"all_tickets_page_{page-1}"))
            nav_row.append(InlineKeyboardButton(f"صفحه {page+1}/{total_pages}", callback_data="ignore"))
            if page < total_pages - 1:
                nav_row.append(InlineKeyboardButton("صفحه بعد ▶️", callback_data=f"all_tickets_page_{page+1}"))
            if nav_row:
                keyboard.append(nav_row)

            # Add refresh button
            keyboard.append([
                InlineKeyboardButton("🔄 به‌روزرسانی", callback_data="refresh_all_tickets")
            ])

            # Add export tickets button
            keyboard.append([
                InlineKeyboardButton("📄 خروجی تیکت‌ها", callback_data="export_all_tickets")
            ])
            
            reply_markup = InlineKeyboardMarkup(keyboard)
            await query.edit_message_text(
                message_text,
                reply_markup=reply_markup,
                parse_mode=ParseMode.MARKDOWN
            )
            
        except Exception as e:
            import telegram
            if isinstance(e, telegram.error.BadRequest) and "Message is not modified" in str(e):
                logger.warning(f"Suppressed Telegram BadRequest: {e}")
                return
            logger.error(f"Error showing all tickets inline: {e}")
            await query.edit_message_text("❌ خطا در نمایش تیکت‌ها.")

    async def paginate_all_tickets_callback(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Handle pagination for all tickets list"""
        query = update.callback_query
        await query.answer()
        if not self._is_admin(query.from_user.id):
            await query.edit_message_text("❌ شما دسترسی لازم برای این عملیات را ندارید.")
            return
        try:
            page = int(query.data.split('_')[-1])
        except ValueError:
            page = 0
        logger.debug(f"paginate_all_tickets_callback: user={query.from_user.id} requested page {page}")
        await self._show_all_tickets_inline(query, page=page)

    async def export_all_tickets_callback(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Generate a Word (.docx) report for *all* tickets and send to admin."""
        query = update.callback_query
        await query.answer()
        if not self._is_admin(query.from_user.id):
            await query.edit_message_text("❌ شما دسترسی لازم برای این عملیات را ندارید.")
            return
        try:
            tickets = self._get_all_tickets()
            if not tickets:
                await query.edit_message_text("📋 هیچ تیکتی یافت نشد.")
                return
            # Convert sqlite3.Row objects to dict for easy access
            tickets = [dict(t) for t in tickets]
            # Sort from oldest to newest
            tickets.sort(key=lambda x: x.get('created_at') or '')
            doc = Document()
            doc.add_heading("گزارش کامل تیکت‌ها", level=1)
            for t in tickets:
                status = t.get('status', '').lower()
                emoji = self._get_status_emoji(status)
                ticket_id = t.get('ticket_id') or t.get('id')
                subject = t.get('subject', '-')
                created_at = t.get('created_at', '-')
                # Heading with emoji and ticket id
                doc.add_heading(f"{emoji} تیکت #{ticket_id} - {subject}", level=2)
                # Status and timestamps
                status_emoji = self._get_status_emoji(status)
                status_readable = str(status).replace('_',' ')
                doc.add_paragraph(f"{status_emoji} وضعیت: {status_readable}")
                doc.add_paragraph(f"📅 تاریخ ثبت: {created_at}")
                # Find answer timestamp & admin reply
                admin_reply = self._get_admin_reply(ticket_id)
                admin_reply_time = None
                try:
                    msgs_all = DatabaseQueries.get_ticket_messages(ticket_id)
                    for m in reversed(msgs_all or []):
                        m = dict(m)
                        if m.get('is_admin'):
                            admin_reply_time = m.get('timestamp')
                            break
                except Exception:
                    pass
                if admin_reply_time:
                    # Find responder name
                    responder_info = self._get_user_info(admin_reply.get('user_id')) if admin_reply else None
                    responder_display = self._format_user_info(responder_info) if responder_info else "نامشخص"
                    doc.add_paragraph(f"📨 تاریخ پاسخ: {admin_reply_time} توسط {responder_display}")
                # User info
                user_info = self._get_user_info(t.get('user_id'))
                user_full = user_info.get('full_name') or f"{user_info.get('first_name','')} {user_info.get('last_name','')}".strip()
                username = user_info.get('username','-')
                phone = user_info.get('phone') or user_info.get('phone_number','-')
                telegram_id = t.get('user_id') or user_info.get('telegram_id') or '-'
                doc.add_paragraph(f"👤 کاربر: {user_full} | @{username} | ID: {telegram_id} | 📞 {phone}")
                # Messages section
                messages = DatabaseQueries.get_ticket_messages(ticket_id)
                if messages:
                    doc.add_heading("پیام‌ها", level=3)
                    for m in messages:
                        m = dict(m)
                        sender = "ادمین" if m.get('is_admin') else "کاربر"
                        ts = m.get('timestamp', '')
                        doc.add_paragraph(f"[{ts}] {sender}: {m.get('message', '')}")
                # Answer status
                if not admin_reply:
                    p = doc.add_paragraph()
                    run = p.add_run("پاسخی داده نشده است")
                    run.font.color.rgb = RGBColor(255,0,0)
                # Separator
                doc.add_paragraph("_"*40)

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            filename = f"tickets_report_{context.bot.id}.docx"
            await context.bot.send_document(
                chat_id=query.from_user.id,
                document=buffer,
                filename=filename,
                caption="📄 گزارش کامل تیکت‌ها",
            )
        except Exception as e:
            logger.error(f"Error exporting tickets: {e}")
            await query.edit_message_text("❌ خطا در تولید خروجی.")

    async def refresh_all_tickets_callback(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Refresh all tickets list"""
        query = update.callback_query
        await query.answer()
        if not self._is_admin(query.from_user.id):
            await query.edit_message_text("❌ شما دسترسی لازم برای این عملیات را ندارید.")
            return
        await self._show_all_tickets_inline(query, page=0)

    @staff_only
    async def show_all_tickets_command(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Send the first page of *all* tickets with inline pagination buttons."""
        try:
            tickets = self._get_all_tickets()
            if not tickets:
                await update.message.reply_text("📋 هیچ تیکتی یافت نشد.")
                return
            # Pagination – page 0
            page = 0
            per_page = 10
            start = page * per_page
            end = start + per_page
            total_pages = (len(tickets) - 1) // per_page + 1
            page_tickets = [dict(t) for t in tickets[start:end]]

            keyboard: list[list[InlineKeyboardButton]] = []
            row: list[InlineKeyboardButton] = []
            message_text = f"📋 لیست تمام تیکت‌ها (صفحه {page+1}/{total_pages}):\n\n"

            for ticket in page_tickets:
                ticket_id = ticket.get('ticket_id') or ticket.get('id')
                user_info = self._get_user_info(ticket.get('user_id'))
                user_display = escape_markdown(self._format_user_info(user_info), version=1)
                subject = escape_markdown(ticket.get('subject', 'بدون موضوع'), version=1)
                created_at = escape_markdown(ticket.get('created_at', ''), version=1)
                status = ticket.get('status', '')
                readable_status = escape_markdown(str(status).replace('_', ' '), version=1)
                emoji = self._get_status_emoji(status)
                message_text += f"{emoji} *تیکت #{ticket_id}* ({readable_status})\n"
                message_text += f"👤 کاربر: {user_display}\n"
                message_text += f"📝 موضوع: {subject}\n"
                message_text += f"📅 تاریخ: {created_at}\n"
                message_text += "───────────────────\n\n"

                row.append(InlineKeyboardButton(f"{emoji} #{ticket_id}", callback_data=f"view_ticket_{ticket_id}"))
                if len(row) == 3:
                    keyboard.append(row)
                    row = []
            if row:
                keyboard.append(row)

            # Navigation buttons
            nav_row: list[InlineKeyboardButton] = []
            if end < len(tickets):
                nav_row.append(InlineKeyboardButton("صفحه بعد ▶️", callback_data="all_tickets_page_1"))
            if nav_row:
                keyboard.append(nav_row)

            keyboard.append([InlineKeyboardButton("🔄 به‌روزرسانی", callback_data="refresh_all_tickets")])
            keyboard.append([InlineKeyboardButton("📄 خروجی تیکت‌ها", callback_data="export_all_tickets")])

            await update.message.reply_text(
                message_text,
                reply_markup=InlineKeyboardMarkup(keyboard),
                parse_mode=ParseMode.MARKDOWN
            )
        except Exception as e:
            logger.error(f"Error showing all tickets (command): {e}")
            await update.message.reply_text("❌ خطا در نمایش تیکت‌ها.")
